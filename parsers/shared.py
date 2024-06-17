#!/usr/bin/env python

from docx import Document
from docx.shared import Pt
#from docx.enum.table import WD_ROW_HEIGHT_RULE
#from docx.enum.style import WD_STYLE_TYPE
import datetime
import yaml
from os import getenv
from glob import iglob
import pandas as pd


lang_pref = getenv('COMPILE_LANGUAGE') or 'language_unknown'
lang_dict = {}
for f in iglob('/content/languages/' + lang_pref + '.y*ml'):
  with open(f,'r') as ymlfile:
    lang_dict = yaml.full_load(ymlfile)

with open('/content/settings.yml', 'r') as f:
  settings = yaml.full_load(f)

def get_val(obj, attr):
  if not obj:
    obj = {attr:attr}
  if lang_pref in obj and attr in obj[lang_pref]:
    value = obj[lang_pref][attr]
  elif attr in obj:
    value = obj[attr]
  elif attr in settings['aggregations']:
    value = []
    for i in settings['aggregations'][attr]:
      indx = str(i)
      tmpval = get_val(obj, indx)
      if tmpval != '':
        value.append(tmpval)
      else:
        value.append(indx) #use it as a string
    value = ''.join(value)
  else:
    return ''

  if isinstance(value, int):
    value = str(value)
  # now we try to convert the value using our language dictionary
  if isinstance(value, str) and value in lang_dict:
    value = lang_dict[value]

  return value

def date2str(inp):
  '''
  Convert 2022-04 to 'April 2022'
  '''
  if isinstance(inp, int):
    return str(inp)
  
  yearstr = inp[:4]
  if len(inp) >= 7:
    monthint = int(inp[5:7])
    monthstr = datetime.date(1900, monthint, 1).strftime('%B')
    return monthstr + ' ' + yearstr
  else:
    return yearstr

def get_max_cols(inp):
  maxcols = 0
  for tpl in inp:
    if len(tpl) > maxcols:
      maxcols = len(tpl)
  return maxcols

# the following function is probably present in pandas
def tuples2monospaced(inp: list, spacing = 1):
  '''
  As input it takes a list of tuples or lists, thus a 2d plane or table
  '''
  inp = [ [str(y) for y in x ] for x in inp ]
  maxcols = get_max_cols(inp)
#
  maxlength_per_col = []
  for i in range(maxcols):
    maxlength_per_col.append(0)
    for tpl in inp:
      if i < len(tpl):
        if len(tpl[i]) > maxlength_per_col[i]:
          maxlength_per_col[i] = len(tpl[i])
#  
  layoutstr = ''
  for i in maxlength_per_col:
    layoutstr += '{:<' + str(i+spacing) + '}'
#  
  result = []
  for tpl in inp:
    line = layoutstr.format(*tpl,'','','','','','','','','') # ugly but it works
    result.append(line)
  return '\n'.join(result)

def tuples2colon(inp: list):
  print('todo, colon separated')

def tuples2docx(inp: list, doc, colswidths = []):
  '''
  The colswidths should be an array of percentages e.g. [25,50,25]
  '''
  table = doc.add_table(rows=0, cols=get_max_cols(inp))
  #table.style = 'TableGrid' #table.add_style('Grid',WD_STYLE_TYPE.TABLE)
  lists = []
  for row in inp:
    tr = table.add_row()
    cells = tr.cells
    for ci in range(len(row)):
      value = row[ci]
      if isinstance(value, list):
        value = '\n'.join(value) #cells[ci].add_paragraph(v, style='List Bullet')
      cells[ci].text = value
    #tr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY #.AUTO
    #tr.height = Pt(12)

  for i in range(len(colswidths)):
    # https://stackoverflow.com/questions/43749177/python-docx-table-column-width
    table.columns[i].width = colswidths[i] * 55000
  return doc

def load_yamls(dir_path):
  objs = []
  for f in iglob(dir_path + '/*.y*ml'):
    with open(f,'r') as ymlfile:
      obj = yaml.full_load(ymlfile)
      if 'end' not in obj:
        obj['end'] = datetime.datetime.now().year
      obj['end'] = str(obj['end']) #make all a str
      if 'archive' not in obj or not obj['archive']:
        objs.append(obj)
      else:
        print('Skipping archived', f)
  
  # sort by end date
  objs.sort(key=lambda obj: obj['end'], reverse=True)
  
  # convert date to long format
  for e in objs:
    if 'end' in e:
      e['end'] = date2str(e['end'])
    if 'start' in e:
      e['start'] = date2str(e['start'])

  return objs

def addHead(doc,title,lvl=1):
  # hack to convert the title to local language
  val = get_val({title:title},title)
  doc.add_heading(val,lvl)

def attribute2tuple(obj,attr):
  value = get_val(obj,attr)
  if value == '':
    return False
  # try to translate the attribute name
  if attr in lang_dict:
    attrname = lang_dict[attr]
  else:
    attrname = attr
  return (attrname,value)

def list2htmllist(arr):
  result = '<ul>'
  for i in arr:
    result += '<li>' + str(i) + '</li>'
  result += '</ul>'
  return result


def tuples2html(inp):
  table = []
  for x in inp:
    row = []
    for y in x:
      if isinstance(y,list):
        y = list2htmllist(y)
      row.append(y)
    table.append(row)

  df = pd.DataFrame(table, dtype=str)

  #pd.set_option('display.max_colwidth', -1)
  #pd.set_option('display.chop_threshold', -1)
  #pd.set_option('display.max_columns', None)
  #pd.set_option('display.expand_frame_repr', False)
  #pd.set_option('max_colwidth', -1)

  #df.apply(lambda x: list2htmllist(x) if isinstance(x,list) else x)
  #df.drop(columns=[0], inplace=True)
  result = df.to_html(bold_rows=False,escape=False,border=0)
  return result

